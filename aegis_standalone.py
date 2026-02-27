#!/usr/bin/env python3
"""
aegis_standalone.py — Project Aegis-CDR Standalone Demo Script

Run this script directly to test PDF and DOCX sanitization
without spinning up the full FastAPI server.

Usage:
    python aegis_standalone.py --file suspicious_doc.pdf
    python aegis_standalone.py --file risky_macro.docx
    python aegis_standalone.py --file evil.exe.pdf  # Will be blocked
    python aegis_standalone.py --demo               # Creates and sanitizes sample docs
"""

import sys
import os
import argparse
import json
import logging
import tempfile
from pathlib import Path

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent))

from utils.validator import SafeTypeValidator
from core.pdf.sanitizer import PDFSanitizer
from core.docx.sanitizer import DocxSanitizer
from core.ai.sentry import AegisSentry

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(name)s] %(levelname)s — %(message)s"
)
log = logging.getLogger("aegis.standalone")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Main Pipeline
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def run_aegis_pipeline(input_file: str, output_dir: str = None, pixel_fallback: bool = False) -> dict:
    """
    Full Aegis CDR pipeline on a single file.
    Returns result dict with threat intelligence.
    """
    input_path = Path(input_file)
    if not input_path.exists():
        raise FileNotFoundError(f"File not found: {input_file}")

    output_dir = Path(output_dir) if output_dir else input_path.parent
    output_path = output_dir / f"SAFE_{input_path.name}"

    print_banner()
    print(f"  📄 Input:  {input_path}")
    print(f"  📦 Output: {output_path}\n")

    # ── Step 1: Magic Number Validation ──────────────────────────────
    print("  [1/4] 🔍 Fingerprinting file type...")
    validator = SafeTypeValidator()

    with open(input_path, "rb") as f:
        raw_bytes = f.read()

    try:
        true_mime = validator.detect_true_type(raw_bytes)
        ext_valid = validator.validate_extension_matches(input_path.name, true_mime)

        print(f"        ✅ True MIME: {true_mime}")
        if not ext_valid:
            print(f"        ⚠️  EXTENSION MISMATCH DETECTED — possible spoofing attempt!")

    except ValueError as e:
        print(f"        🚨 BLOCKED: {e}")
        return {"status": "blocked", "reason": str(e)}

    # ── Step 2-4: Decompose → Disarm → Reconstruct ───────────────────
    print(f"  [2/4] 💣 Decomposing and disarming...")

    result = {}
    if true_mime == "application/pdf":
        sanitizer = PDFSanitizer(str(input_path), str(output_path), pixel_fallback=pixel_fallback)
        result = sanitizer.sanitize()
    elif "wordprocessingml" in true_mime:
        sanitizer = DocxSanitizer(str(input_path), str(output_path))
        result = sanitizer.sanitize()
    else:
        print(f"  ❌ Unsupported format: {true_mime}")
        return {"status": "unsupported", "mime": true_mime}

    items_removed = result.get("items_removed", [])
    print(f"        Found {len(items_removed)} threat(s):")
    for item in items_removed:
        print(f"        ├── ⚡ {item}")
    if not items_removed:
        print("        └── ✅ No active threats detected")

    # ── AI Sentry Analysis ───────────────────────────────────────────
    print(f"\n  [3/4] 🤖 AI Sentry analysis...")
    sentry = AegisSentry()
    summary = sentry.summarize(
        items_removed,
        result.get("page_count_original", 0),
        result.get("page_count_sanitized", 0)
    )
    risk = sentry.risk_score(items_removed)

    print(f"        Risk Level : {risk['level']} ({risk['score']}/100)")
    print(f"        Rationale  : {risk['rationale']}")

    # ── Output Summary ────────────────────────────────────────────────
    print(f"\n  [4/4] 📋 Reconstruction complete.")
    print(f"\n{'─'*60}")
    print(f"  AEGIS THREAT REPORT")
    print(f"{'─'*60}")
    print(f"  {summary}")
    if result.get("fallback_used"):
        print(f"\n  ⚠️  PIXEL FALLBACK ENGAGED — document was rasterized")
    print(f"{'─'*60}\n")

    final_result = {
        "status": "sanitized",
        "input": str(input_path),
        "output": str(output_path),
        "true_mime": true_mime,
        "items_removed": items_removed,
        "threat_summary": summary,
        "risk": risk,
        "page_count_original": result.get("page_count_original", 0),
        "page_count_sanitized": result.get("page_count_sanitized", 0),
        "fallback_used": result.get("fallback_used", False),
    }

    # Save JSON report
    report_path = output_dir / f"aegis_report_{input_path.stem}.json"
    with open(report_path, "w") as f:
        json.dump(final_result, f, indent=2)
    print(f"  📊 Report saved: {report_path}\n")

    return final_result


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Demo Mode: Create Sample Files for Testing
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def create_demo_pdf(output_path: str):
    """Create a sample PDF with embedded JavaScript for demo purposes."""
    try:
        import fitz
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 100), "Project Aegis Demo Document", fontsize=18)
        page.insert_text((72, 140), "This PDF contains simulated active content.", fontsize=11)
        page.insert_text((72, 160), "Aegis will detect and remove it during sanitization.", fontsize=11)

        # Inject a JavaScript alert (simulated threat)
        # Note: In a real malicious PDF this would be in the catalog
        page.insert_text((72, 220), "Hidden: /JavaScript << /S /JavaScript /JS (app.alert('XSS');) >>", fontsize=8)
        page.insert_text((72, 240), "Hidden: /OpenAction << /S /Launch /Win << /F (cmd.exe) >> >>", fontsize=8)

        doc.save(output_path)
        doc.close()
        print(f"  ✅ Demo PDF created: {output_path}")
    except ImportError:
        print("  ⚠️  PyMuPDF not installed. Install with: pip install PyMuPDF")


def create_demo_docx(output_path: str):
    """Create a sample DOCX for demo purposes."""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Project Aegis Demo Document", 0)
        doc.add_paragraph("This document demonstrates Aegis CDR sanitization.")
        doc.add_paragraph("In a real malicious DOCX, macros and external links would be present.")
        doc.save(output_path)
        print(f"  ✅ Demo DOCX created: {output_path}")
    except ImportError:
        print("  ⚠️  python-docx not installed. Install with: pip install python-docx")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Banner & CLI
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def print_banner():
    print("""
╔══════════════════════════════════════════════════════════════╗
║           PROJECT AEGIS — Content Disarm & Reconstruction    ║
║           AI-Powered Document Sanitization Engine v1.0       ║
╚══════════════════════════════════════════════════════════════╝
""")


def main():
    parser = argparse.ArgumentParser(description="Project Aegis-CDR Standalone Runner")
    parser.add_argument("--file", "-f", help="Path to file to sanitize")
    parser.add_argument("--output-dir", "-o", help="Output directory (default: same as input)")
    parser.add_argument("--pixel-fallback", action="store_true",
                        help="Enable pixel-only fallback for high-risk PDFs")
    parser.add_argument("--demo", action="store_true",
                        help="Create and sanitize sample demo files")
    parser.add_argument("--json", action="store_true",
                        help="Output results as JSON to stdout")

    args = parser.parse_args()

    if args.demo:
        print_banner()
        print("  Creating demo files...\n")
        tmp = tempfile.mkdtemp()
        pdf_path = os.path.join(tmp, "demo_threat.pdf")
        docx_path = os.path.join(tmp, "demo_macro.docx")

        create_demo_pdf(pdf_path)
        create_demo_docx(docx_path)

        print("\n  Running PDF sanitization demo...\n")
        run_aegis_pipeline(pdf_path, tmp)

        print("\n  Running DOCX sanitization demo...\n")
        run_aegis_pipeline(docx_path, tmp)

    elif args.file:
        result = run_aegis_pipeline(
            args.file,
            args.output_dir,
            pixel_fallback=args.pixel_fallback
        )
        if args.json:
            print(json.dumps(result, indent=2))
    else:
        parser.print_help()


if __name__ == "__main__":
    main()